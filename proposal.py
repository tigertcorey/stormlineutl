"""
DOCX proposal generator for Stormline Utilities.
Fills STORMLINE_MASTER_PROPOSAL_v5.docx template with estimate data.
No unit prices shown to GC — descriptions, quantities, and section subtotals only.
"""

import os
import re
import copy
from datetime import datetime
from lxml import etree

TEMPLATE_PATH = "/mnt/c/Users/Corey Tigert/OneDrive/Desktop/STORMLINE_MASTER_PROPOSAL_v5.docx"
PROJECTS_DIR = "/mnt/c/Users/Corey Tigert/OneDrive/Desktop/PROJECTS"

# Template table index for each utility type
SECTION_TABLE = {"storm": 1, "water": 2, "sanitary": 3, "sewer": 3, "fire": 4, "fdc": 4}
DATA_ROW_START = {1: 1, 2: 1, 3: 1, 4: 0}   # table 4 has no header row

# Item-name keywords that override section-based routing
_ITEM_WATER    = ("hydrant", "valve", "gate valve", "dcva", "backflow", "meter", "service tap",
                  "fire hydrant", "detector check", "rpz", "tapping sleeve")
_ITEM_SEWER    = ("cleanout", "drop mh", "connect to existing mh", "service connection")
_ITEM_FIRE     = ("fire line", "fdc", "fire service", "detector check valve", "6 di", "8 di")
_ITEM_STORM    = ("rcp", "hdpe storm", "curb inlet", "area inlet", "junction box",
                  "catch basin", "storm mh", "headwall")


def _classify_item(item_name: str, section_name: str) -> str:
    """Return utility group based on item name + section, most specific wins."""
    n = item_name.lower()
    s = section_name.lower()

    if any(k in n for k in _ITEM_FIRE):
        return "fire"
    if any(k in n for k in _ITEM_WATER):
        return "water"
    if any(k in n for k in _ITEM_SEWER):
        return "sanitary"
    if any(k in n for k in _ITEM_STORM):
        return "storm"

    # Fall back to section label
    if any(w in s for w in ("fire", "fdc")):
        return "fire"
    if "water" in s:
        return "water"
    if any(w in s for w in ("sanitary", "sewer")):
        return "sanitary"
    return "storm"


def _set_cell(cell, text: str):
    """Set cell text, preserving first run's font/bold from template."""
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = text
        for extra in para.runs[1:]:
            extra.text = ""
    else:
        para.add_run(text)


def _copy_row(table, src_row_idx: int):
    """Duplicate a row in a table (deep copy of XML) and insert before subtotal row."""
    from docx.oxml.ns import qn
    src_tr = table.rows[src_row_idx]._tr
    new_tr = copy.deepcopy(src_tr)
    # Insert before last row
    table._tbl.insert(len(table._tbl) - 1, new_tr)
    return table.rows[-2]   # newly inserted row is now second-to-last


def _fill_section(table, tbl_idx: int, items: list, subtotal: float):
    """
    Fill a scope section table with all items, adding rows if needed.
    items: list of line_item dicts from estimate_from_takeoff.
    """
    start = DATA_ROW_START[tbl_idx]
    template_slots = 6  # data rows in the original template

    # Expand table if more items than slots
    extra_needed = max(0, len(items) - template_slots)
    for _ in range(extra_needed):
        _copy_row(table, start)  # clone first data row

    # Fill all data rows
    for slot in range(max(template_slots, len(items))):
        # Row index: start + slot, but after adding rows the layout may shift
        # Safer: use start + slot, capped to available rows before subtotal
        row_idx = start + slot
        # Rows: 0..start-1 = header; start..(-2 or -1) = data; last = subtotal
        if tbl_idx == 4:
            last_data_row = len(table.rows) - 3   # table 4 has subtotal + total bid
        else:
            last_data_row = len(table.rows) - 2   # table 1-3: header + data + subtotal
        if row_idx > last_data_row:
            break

        row = table.rows[row_idx]
        if slot < len(items):
            li = items[slot]
            qty = li.get("qty", 0)
            qty_str = f"{int(qty):,}" if qty == int(qty) else f"{qty:,.1f}"
            unit_cost = li.get("unit_cost", 0)
            extension = li.get("extension", 0)
            _set_cell(row.cells[0], str(slot + 1))
            _set_cell(row.cells[1], li.get("item", ""))
            _set_cell(row.cells[2], li.get("unit", "LF"))
            _set_cell(row.cells[3], qty_str)
            _set_cell(row.cells[4], f"${unit_cost:,.2f}")
            _set_cell(row.cells[5], f"${extension:,.2f}")
        else:
            for c in row.cells:
                _set_cell(c, "")

    # Subtotal cell — second-to-last row for table 4, last row for 1-3
    subtotal_row = table.rows[-2] if tbl_idx == 4 else table.rows[-1]
    _set_cell(subtotal_row.cells[5], f"${subtotal:,.2f}")


def generate_docx(estimate_data: dict, project_address: str = "",
                  gc_contact: str = "", engineer: str = "") -> dict:
    """
    Fill the v5 proposal template and save to Desktop.
    Returns {"success": True, "path": "..."} or {"success": False, "error": "..."}.
    estimate_data: full result or 'data' dict from estimate_from_takeoff.
    """
    try:
        from docx import Document
    except ImportError:
        return {"success": False, "error": "python-docx not installed"}

    if not os.path.exists(TEMPLATE_PATH):
        return {"success": False, "error": f"Template not found: {TEMPLATE_PATH}"}

    d = estimate_data.get("data", estimate_data)
    job_name   = d.get("job_name", "Project")
    gc_name    = d.get("gc_name", "")
    total_bid  = d.get("total_bid", 0)
    line_items = d.get("line_items", [])
    mob        = d.get("mobilization", 0)
    testing    = d.get("testing", 0)
    sales_tax  = d.get("sales_tax", 0)
    today      = datetime.now().strftime("%B %d, %Y")

    doc = Document(TEMPLATE_PATH)

    # ── Table 0: Job header ───────────────────────────────────────────────
    field_map = {
        "JOB NAME:":       job_name,
        "DATE:":           today,
        "CITY:":           project_address,
        "GC / OWNER:":     f"{gc_name}  {gc_contact}".strip(),
        "CIVIL ENGINEER:": engineer,
    }
    for row in doc.tables[0].rows:
        key = row.cells[0].text.strip()
        if key in field_map:
            _set_cell(row.cells[1], field_map[key])

    # ── Group items by utility type (name-aware routing) ─────────────────
    groups: dict[str, list]  = {"storm": [], "water": [], "sanitary": [], "fire": []}
    subtotals: dict[str, float] = {"storm": 0.0, "water": 0.0, "sanitary": 0.0, "fire": 0.0}

    for li in line_items:
        g = _classify_item(li.get("item", ""), li.get("section", ""))
        groups[g].append(li)
        subtotals[g] += li.get("extension", 0.0)

    # ── Fill scope tables ─────────────────────────────────────────────────
    for util_type, items in groups.items():
        tbl_idx = SECTION_TABLE[util_type]
        _fill_section(doc.tables[tbl_idx], tbl_idx, items, subtotals[util_type])

    # ── Table 4 last row: total bid with breakdown ────────────────────────
    total_row = doc.tables[4].rows[-1]
    breakdown = (
        f"Storm ${subtotals['storm']:,.0f}  |  Water ${subtotals['water']:,.0f}  |  "
        f"Sewer ${subtotals['sanitary']:,.0f}  |  Fire ${subtotals['fire']:,.0f}  |  "
        f"Mob ${mob:,.0f}  |  Testing ${testing:,.0f}  |  Tax ${sales_tax:,.0f}"
    )
    _set_cell(total_row.cells[0], "TOTAL BASE BID:")
    for c in total_row.cells[1:-1]:
        _set_cell(c, breakdown)
    _set_cell(total_row.cells[-1], f"${total_bid:,.2f}")

    # ── Save into project folder ──────────────────────────────────────────
    safe_name = re.sub(r"[^\w\s-]", "", job_name).strip().replace(" ", "_")
    # Try to match an existing folder (case-insensitive) before creating a new one
    project_dir = None
    if os.path.isdir(PROJECTS_DIR):
        job_upper = job_name.upper()
        for folder in os.listdir(PROJECTS_DIR):
            if folder.upper() == job_upper or folder.upper() == safe_name.upper().replace("_", " "):
                project_dir = os.path.join(PROJECTS_DIR, folder)
                break
    if project_dir is None:
        project_dir = os.path.join(PROJECTS_DIR, safe_name)
    os.makedirs(project_dir, exist_ok=True)
    out_path = os.path.join(project_dir, f"PROPOSAL_{safe_name}_{datetime.now().strftime('%Y%m%d')}.docx")
    doc.save(out_path)
    return {"success": True, "path": out_path, "total_bid": total_bid}
